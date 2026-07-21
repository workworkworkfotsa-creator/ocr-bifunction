"""Stage ① LIRE — read the text out of a document.

The entry stage is NOT "OCR": it is "get the text out of a document by the means
fit for its type". OCR is only ONE backend, for image-only content. Born-digital
PDFs and .docx carry their text natively and never touch an OCR engine.

Every backend yields TextLine geometry, not just a flat string: the bounding box
is the spatial anchor stage ③ rebuilds fields from (e.g. the value to the right of
"NOM / Surname" → the surname). Raw words carry no links; their positions do.

The OcrEngine Protocol is the jettisonable slot from the cadrage: RapidOCR,
Tesseract or granite-docling all plug in behind it, interchangeable. Which engine
gets selected is a function of the regime (API fast-path vs backoffice batch) and
of confidence — that selection is the "pont" of the dual model, decided upstream,
not here.
"""

from __future__ import annotations

import time
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Protocol, runtime_checkable

from ocr_bifunction.conversion_guard import page_count
from ocr_bifunction.resilient_conversion import (
    PageRangeConverter,
    reconcile_page_range_conversion,
)
from ocr_bifunction.text_integrity_guard import (
    TextIntegrityAssessment,
    assess_text_integrity,
)

# A factory that binds a per-document `PageRangeConverter` (e.g. Docling bound to one path, sharing
# one loaded model): the heavy resilient read needs the path both to count pages and to convert.
HeavyPageConverterFactory = Callable[[Path], PageRangeConverter]

# A PDF page yielding fewer than this many characters of native text is treated as
# image-only (scanned) and routed to the OCR engine instead of the text layer.
TEXT_LAYER_MINIMUM_CHARACTERS = 10

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}


@dataclass
class TextLine:
    """One text region with its geometry — the spatial anchor templates rely on.

    `bbox` is axis-aligned `(x0, y0, x1, y1)` in a TOP-LEFT-origin page/image coordinate
    system, with `x0 <= x1` and `y0 <= y1`: every engine normalizes to this (RapidOCR/Docling
    via min/max, the PDF text layer via PyMuPDF, the VLM via a synthetic top-to-bottom box) so
    the template "value below / value to the right" geometry has ONE orientation. `confidence`
    is None for a native text-layer region (exact — trust it) or a float in `[0, 1]` for an OCR
    score (the "douteux → humain" signal).

    The invariants are enforced at construction (fail-loud): a malformed engine output — a
    wrong-length box, a bottom-left origin left un-flipped, a score outside [0, 1] — surfaces
    HERE with a clear error, not as silent geometry garbage in `_value_below` downstream. That
    is the hardening of the jettisonable-engine contract: the slot may be swapped, its OUTPUT
    shape may not drift. Read the geometry through the named accessors (`x0`/`y0`/`x1`/`y1`/
    `width`/`height`), not by indexing `bbox` — the box is the contract, the tuple is storage.

    `page_width`/`page_height` are the size of the page this line was read from, IN THE SAME
    UNIT as `bbox` — PDF points for a text-layer read, pixels for an OCR'd render. They exist
    because a box alone is NOT placeable: the same tuple means different things at 72 dpi and
    at 200 dpi, and a consumer wanting to draw the region has no way to tell. `0.0` means
    UNKNOWN, and unknown is a real state (the VLM lane's boxes are synthetic reading-order
    scaffolding, not positions) — see `ProvenanceSpan.from_line`, which refuses to mint
    provenance without them rather than inventing a frame of reference.
    """

    text: str
    bbox: tuple[float, float, float, float]
    confidence: float | None = None
    page_index: int = 0
    page_width: float = 0.0
    page_height: float = 0.0
    # Per-word geometry INSIDE this line, when the backend can give it (the PDF text layer can;
    # OCR engines already emit line-sized boxes and leave this empty). Lets a consumer narrow a
    # box to the words a value actually occupies — see `WordSpan`.
    word_spans: list[WordSpan] = field(default_factory=list)

    def __post_init__(self) -> None:
        if self.page_width < 0 or self.page_height < 0:
            raise ValueError(
                "TextLine page dimensions must be non-negative (0 = unknown), got "
                f"{self.page_width!r} x {self.page_height!r}"
            )
        if len(self.bbox) != 4:
            raise ValueError(
                f"TextLine.bbox must be (x0, y0, x1, y1), got {self.bbox!r}"
            )
        x0, y0, x1, y1 = self.bbox
        if x1 < x0 or y1 < y0:
            raise ValueError(
                "TextLine.bbox must be axis-aligned, top-left origin (x0<=x1, y0<=y1), "
                f"got {self.bbox!r}"
            )
        if self.confidence is not None and not 0.0 <= self.confidence <= 1.0:
            raise ValueError(
                f"TextLine.confidence must be None or within [0, 1], got {self.confidence!r}"
            )

    @property
    def x0(self) -> float:
        return self.bbox[0]

    @property
    def y0(self) -> float:
        return self.bbox[1]

    @property
    def x1(self) -> float:
        return self.bbox[2]

    @property
    def y1(self) -> float:
        return self.bbox[3]

    @property
    def width(self) -> float:
        return self.bbox[2] - self.bbox[0]

    @property
    def height(self) -> float:
        return self.bbox[3] - self.bbox[1]


@dataclass
class WordSpan:
    """One word of a line: WHERE it sits in the line's text, and where it sits on the page.

    `start`/`end` are character offsets into the owning `TextLine.text`; `bbox` is in that
    line's native units. This exists so a value found at known CHARACTER positions can be
    narrowed to the box of just those words — a born-digital block is paragraph-sized, so
    highlighting the whole block tells a reviewer "somewhere in here" instead of "there".

    Offsets, never spelling: the same word occurs many times on a page, so looking one up by
    its text lands on an arbitrary occurrence. Measured 2026-07-21 on a real invoice — matching
    by text made the box 3x LARGER than the block it was meant to shrink, because the words of
    a date also appear in the document title.
    """

    start: int
    end: int
    bbox: tuple[float, float, float, float]


@dataclass
class ProvenanceSpan:
    """Where one piece of a downstream product came from in the SOURCE document: page + box.

    This is the link from read text back to the original document — so a retrieved passage OR
    an extracted field can be shown verbatim AND located for the human to verify against the
    source.

    `bbox` is `(x0, y0, x1, y1)` NORMALIZED to the page: each value is a FRACTION in `[0, 1]`
    of the page's width or height. Deliberately not the reader's native units, because those
    differ by backend — PDF points at 72 dpi for a text layer, pixels of a 200 dpi render for
    OCR — and a raw tuple is therefore unplaceable without knowing which. Normalized, a
    consumer draws the region with no unit, no dpi and no page size to carry, and the value
    survives a change of render resolution or OCR engine.

    It lives HERE, next to `TextLine`, because it is a projection of one — not a RAG concept:
    the retrieval lane (`rag.Chunk.spans`) and the structured lane (`template.ExtractedField
    .spans`) both need the same vocabulary, and neither should import the other.
    """

    page_index: int
    bbox: tuple[float, float, float, float]

    @classmethod
    def from_line(
        cls,
        line: TextLine,
        *,
        bbox: tuple[float, float, float, float] | None = None,
    ) -> ProvenanceSpan | None:
        """The span of a read line, normalized — the only way provenance is created.

        `bbox` narrows the span to a REGION of the line (the words a value occupies) while
        reusing the line's page frame; None means the whole line, which is all a backend
        without per-word geometry can offer.

        Returns None when the line carries no page dimensions. That is not an error to swallow
        but a REAL state: the VLM lane emits synthetic reading-order boxes that encode no
        position, so there is no frame of reference to normalize against. Minting a span from
        one would fabricate a location — the thing this codebase refuses to do. Callers treat
        None as "no provenance", which is exactly what it is.
        """
        if line.page_width <= 0 or line.page_height <= 0:
            return None
        x0, y0, x1, y1 = bbox if bbox is not None else line.bbox
        return cls(
            line.page_index,
            (
                x0 / line.page_width,
                y0 / line.page_height,
                x1 / line.page_width,
                y1 / line.page_height,
            ),
        )


@dataclass
class ReadResult:
    """What stage ① hands to stage ② CATÉGORISER.

    `confidence` is the legibility signal the confidence gate routes on:
    None  -> backend gives no score (native text layer = exact, trust it);
    float -> mean OCR score in [0, 1]; a low value is the "douteux → humain" signal.

    `lines` carries the geometry: stage ③ rebuilds "Nom: …, Prénom: …" from the
    boxes, not from the flat `text` (which is only the lines joined for display).
    """

    document_path: Path
    backend_name: str
    text: str = ""
    lines: list[TextLine] = field(default_factory=list)
    confidence: float | None = None
    page_count: int = 0
    character_count: int = 0
    elapsed_seconds: float = 0.0
    needs_ocr: bool = False  # routed to an OCR engine, but none was wired in
    error: str | None = None
    # Heavy resilient read only (empty otherwise): pages never produced even after the smallest
    # batch (genuinely bad -> human review) and pages present but below the layout threshold. The
    # page-grain twin of the CI `missing` signal — an incomplete read never passes for a whole one.
    missing_pages: list[int] = field(default_factory=list)
    low_form_pages: list[int] = field(default_factory=list)
    # Character-integrity signal, None when there was no text to assess. The extracted characters
    # can be MOJIBAKE on a perfectly native PDF (a subset font with a broken ToUnicode CMap) — a
    # corruption every text-layer extractor inherits identically, so it is checked HERE, once, on
    # the output, rather than per-engine. Orthogonal to `confidence` (legibility of a recognition)
    # and to `missing_pages` (completeness): those can all be perfect while the characters are wrong.
    text_integrity: TextIntegrityAssessment | None = None


@runtime_checkable
class OcrEngine(Protocol):
    """The jettisonable OCR slot. Any engine that turns a rendered image into lines fits.

    Output contract (enforced by `TextLine`, so a violating engine fails loud): `recognize`
    returns `TextLine`s whose `bbox` is axis-aligned, top-left origin (x0<=x1, y0<=y1) and
    whose `confidence` is None (no per-line score) or a float in `[0, 1]`. The engine may be
    swapped freely (RapidOCR / Docling / a VLM); its output SHAPE may not drift.
    """

    name: str

    def recognize(self, image_png_bytes: bytes) -> list[TextLine]:
        """Return the recognized text lines (with geometry) for one rendered PNG image."""
        ...


def read_document(
    document_path: Path,
    ocr_engine: OcrEngine | None = None,
    *,
    heavy_page_converter_factory: HeavyPageConverterFactory | None = None,
) -> ReadResult:
    """Route a document to the backend fit for its type and return its text.

    `heavy_page_converter_factory` (opt-in) selects the RESILIENT read for PDFs — split into
    page-range batches, retry any dropped page on a smaller batch, reconcile (for the RAG/contract
    lane where a heavy converter's reading-order is the value and mid-document memory drops are the
    risk). Default None keeps the existing per-page light read untouched, so every current caller is
    unaffected."""
    suffix = document_path.suffix.lower()
    if suffix == ".pdf":
        if heavy_page_converter_factory is not None:
            result = _read_pdf_resilient(document_path, heavy_page_converter_factory)
        else:
            result = _read_pdf(document_path, ocr_engine)
    elif suffix == ".docx":
        result = _read_docx(document_path)
    elif suffix in IMAGE_SUFFIXES:
        result = _read_image(document_path, ocr_engine)
    else:
        result = ReadResult(
            document_path=document_path,
            backend_name="(unsupported)",
            error=f"unsupported file type: {suffix}",
        )

    # ONE seam, every backend: whatever produced the characters (PDF text layer, OCR engine,
    # docx, heavy resilient converter), they are checked for corruption here. That is what makes
    # the guard model-agnostic — the failure it catches belongs to the SOURCE, not to any engine.
    # Left None when there is nothing to assess (an image-only page with no OCR engine wired):
    # "not assessed" is honest, "clean" on empty text would not be.
    if result.text.strip():
        result.text_integrity = assess_text_integrity(result.text)
    return result


def _read_pdf_resilient(
    document_path: Path,
    heavy_page_converter_factory: HeavyPageConverterFactory,
) -> ReadResult:
    """Heavy multi-page read that survives a converter dropping pages mid-document.

    Splits the PDF into page-range batches, retries any dropped page under a decreasing batch-size
    schedule, and reconciles the produced pages into one document ordered by page number. The rich
    per-page markdown (reading-order preserved) is joined into `text` for the RAG chunker; the
    completeness verdict rides on `missing_pages`/`low_form_pages` so the router routes an incomplete
    read to a human.

    Per-line geometry IS carried (since 2026-07-21): the converter knows where each piece of text sat
    and that is the only moment it is known, so the adapter hands it over as `TextSpan`s and they are
    rebuilt into `TextLine`s here. That closes the provenance chain on the heavy lane — `rag.py`
    already packs chunks WITH their page+bbox spans, so a passage read this way can now be shown back
    to a human at its place in the original. Empty when a converter exposes no geometry: absent
    provenance stays absent, never fabricated."""
    started_at = time.perf_counter()
    expected_page_count = page_count(document_path)
    page_range_converter = heavy_page_converter_factory(document_path)
    conversion = reconcile_page_range_conversion(
        expected_page_count, page_range_converter
    )
    combined_text = "\n\n".join(page.markdown for page in conversion.page_results)
    # `page_number` is absolute 1-based; `TextLine.page_index` is 0-based (reader contract).
    lines = [
        TextLine(
            span.text,
            span.bbox,
            None,
            page.page_number - 1,
            page_width=page.page_width,
            page_height=page.page_height,
        )
        for page in conversion.page_results
        for span in page.text_spans
    ]
    return ReadResult(
        document_path=document_path,
        backend_name="resilient-page-range",
        text=combined_text,
        lines=lines,
        confidence=None,
        page_count=expected_page_count,
        character_count=len(combined_text),
        elapsed_seconds=time.perf_counter() - started_at,
        missing_pages=conversion.missing_page_numbers,
        low_form_pages=conversion.low_form_page_numbers,
    )


def _mean_confidence(lines: list[TextLine]) -> float | None:
    scores = [line.confidence for line in lines if line.confidence is not None]
    return sum(scores) / len(scores) if scores else None


def _word_spans_in_block(block_text: str, words: list[tuple]) -> list[WordSpan]:
    """Locate each word INSIDE the block's text by walking a cursor forward.

    Sequential on purpose, never `block_text.find(word)` from the start: the same word occurs
    repeatedly on a page and inside a block, so a spelling lookup lands on whichever occurrence
    comes first — which is usually not the one being located. Measured on a real invoice: text
    matching produced a box 3x LARGER than the block it was supposed to shrink, because a
    date's words also appeared in the document title. Advancing the cursor keeps every word at
    its own place, since PyMuPDF yields words in reading order within a block.

    A word the block text does not contain verbatim (ligature, hyphenation, an exotic space) is
    SKIPPED rather than placed approximately: one missing word narrows the box slightly less,
    while a mis-placed one would point somewhere wrong.
    """
    spans: list[WordSpan] = []
    cursor = 0
    for word in words:
        x0, y0, x1, y1, text = word[0], word[1], word[2], word[3], word[4]
        if not text:
            continue
        found = block_text.find(text, cursor)
        if found < 0:
            continue
        spans.append(WordSpan(found, found + len(text), (x0, y0, x1, y1)))
        cursor = found + len(text)
    return spans


def _read_pdf(document_path: Path, ocr_engine: OcrEngine | None) -> ReadResult:
    import pymupdf

    started_at = time.perf_counter()
    page_texts: list[str] = []
    lines: list[TextLine] = []
    ocr_page_count = 0
    has_text_layer = False
    needs_ocr = False

    with pymupdf.open(document_path) as document:
        page_count = document.page_count
        for page_index, page in enumerate(document):
            native_text = page.get_text().strip()
            if len(native_text) >= TEXT_LAYER_MINIMUM_CHARACTERS:
                has_text_layer = True
                page_texts.append(native_text)
                # Same page, word grain: PyMuPDF numbers words by the block they belong to,
                # so each block gets exactly its own words (word tuple = x0,y0,x1,y1,text,
                # block_no,line_no,word_no).
                words_by_block: dict[int, list[tuple]] = {}
                for word in page.get_text("words"):
                    words_by_block.setdefault(word[5], []).append(word)
                for x0, y0, x1, y1, block_text, block_no, block_type in page.get_text(
                    "blocks"
                ):
                    cleaned = block_text.strip()
                    if block_type == 0 and cleaned:
                        # Block coordinates are PDF points; page.rect is the matching frame.
                        lines.append(
                            TextLine(
                                cleaned,
                                (x0, y0, x1, y1),
                                None,
                                page_index,
                                page_width=page.rect.width,
                                page_height=page.rect.height,
                                word_spans=_word_spans_in_block(
                                    cleaned, words_by_block.get(block_no, [])
                                ),
                            )
                        )
                continue
            # Image-only page: render it and hand it to the OCR engine, if any.
            if ocr_engine is None:
                needs_ocr = True
                continue
            pixmap = page.get_pixmap(dpi=200)
            page_lines = ocr_engine.recognize(pixmap.tobytes("png"))
            for line in page_lines:
                line.page_index = page_index
            lines.extend(page_lines)
            page_texts.append("\n".join(line.text for line in page_lines))
            ocr_page_count += 1

    combined_text = "\n".join(page_texts)
    backend_name = "pymupdf-text"
    if ocr_page_count:
        backend_name = (
            f"pymupdf+{ocr_engine.name}" if has_text_layer else ocr_engine.name
        )
    return ReadResult(
        document_path=document_path,
        backend_name=backend_name,
        text=combined_text,
        lines=lines,
        confidence=_mean_confidence(lines),
        page_count=page_count,
        character_count=len(combined_text),
        elapsed_seconds=time.perf_counter() - started_at,
        needs_ocr=needs_ocr,
    )


def _read_docx(document_path: Path) -> ReadResult:
    from docx import Document

    started_at = time.perf_counter()
    document = Document(str(document_path))
    blocks = [
        paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    combined_text = "\n".join(blocks)
    # .docx is reflowable: it has no page geometry, so no TextLine boxes here.
    return ReadResult(
        document_path=document_path,
        backend_name="python-docx",
        text=combined_text,
        confidence=None,
        page_count=1,
        character_count=len(combined_text),
        elapsed_seconds=time.perf_counter() - started_at,
    )


def _read_image(document_path: Path, ocr_engine: OcrEngine | None) -> ReadResult:
    if ocr_engine is None:
        return ReadResult(
            document_path=document_path,
            backend_name="(ocr pending)",
            needs_ocr=True,
        )
    started_at = time.perf_counter()
    lines = ocr_engine.recognize(document_path.read_bytes())
    combined_text = "\n".join(line.text for line in lines)
    return ReadResult(
        document_path=document_path,
        backend_name=ocr_engine.name,
        text=combined_text,
        lines=lines,
        confidence=_mean_confidence(lines),
        page_count=1,
        character_count=len(combined_text),
        elapsed_seconds=time.perf_counter() - started_at,
    )
